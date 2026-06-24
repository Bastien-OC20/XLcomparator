"""
Comparison engine for XLcomparator.
Compares two Excel workbooks cell-by-cell and returns a list
of differences.
"""

from __future__ import annotations

import io
import xml.etree.ElementTree as ET
from collections import Counter
from dataclasses import dataclass
from typing import Dict, List, Any, NamedTuple, Optional
from xml.dom import minidom

import numpy as np
import pandas as pd
from openpyxl.utils import (
    get_column_letter,
    column_index_from_string,
)


# ---------------------------------------------------------------------------
# Options & result types
# ---------------------------------------------------------------------------

@dataclass
class CompareOptions:
    """Options controlling how the comparison is performed."""
    key_col: Optional[str] = None
    ignore_case: bool = False
    ignore_whitespace: bool = False
    normalize_types: bool = True


class CompareResult(NamedTuple):
    """Result returned by compare_workbooks."""
    differences: List[Dict[str, Any]]
    sheet_stats: Dict[str, int]


# ---------------------------------------------------------------------------
# File reading
# ---------------------------------------------------------------------------

def _read_workbook(file_obj) -> Dict[str, pd.DataFrame]:
    """Read all sheets into a dict of DataFrames (values as str)."""
    file_obj.seek(0)
    xl = pd.ExcelFile(file_obj)
    sheets: Dict[str, pd.DataFrame] = {}
    for name in xl.sheet_names:
        df = xl.parse(name, header=None, dtype=str)
        df = df.fillna("")
        df.columns = list(range(len(df.columns)))
        sheets[str(name)] = df
    return sheets


# ---------------------------------------------------------------------------
# Normalization
# ---------------------------------------------------------------------------

def _normalize_col(
    series: pd.Series,
    opts: CompareOptions,
) -> pd.Series:
    """Normalize a column Series for comparison (vectorized).

    Only applies numeric/date coercion when ALL non-empty values
    in the column can be consistently parsed, avoiding false
    positives on mixed-type columns.
    """
    s = series.astype(str).fillna("")
    if opts.ignore_whitespace:
        s = s.str.strip()
    if opts.ignore_case:
        s = s.str.lower()
    if not opts.normalize_types:
        return s

    non_empty = s[s != ""]
    if len(non_empty) == 0:
        return s

    # Numeric — only if every non-empty value parses.
    numeric_test = pd.to_numeric(non_empty, errors="coerce")
    if numeric_test.notna().all():
        numeric = pd.to_numeric(
            s.replace("", np.nan), errors="coerce"
        )

        def _fmt(v):
            if pd.isna(v):
                return ""
            try:
                if v == int(v):
                    return str(int(v))
            except (ValueError, OverflowError):
                pass
            return str(round(v, 10))

        return numeric.map(_fmt)

    # Datetime — only if every non-empty value parses.
    try:
        date_test = pd.to_datetime(
            non_empty, dayfirst=True, errors="coerce"
        )
        if date_test.notna().all():
            dates = pd.to_datetime(
                s.where(s != "", other=None),
                dayfirst=True,
                errors="coerce",
            )
            return dates.dt.strftime("%Y-%m-%d").fillna("")
    except Exception:
        pass

    return s


# ---------------------------------------------------------------------------
# Positional comparison (vectorized)
# ---------------------------------------------------------------------------

def _compare_positional(
    ref_df: pd.DataFrame,
    cmp_df: pd.DataFrame,
    sheet: str,
    opts: CompareOptions,
) -> tuple[List[Dict[str, Any]], int]:
    """Vectorized positional comparison using np.where.

    Returns (differences, total_cells_compared).
    """
    max_rows = max(len(ref_df), len(cmp_df))
    max_cols = max(
        ref_df.shape[1] if not ref_df.empty else 0,
        cmp_df.shape[1] if not cmp_df.empty else 0,
    )
    if max_rows == 0 or max_cols == 0:
        return [], 0

    ref_df = ref_df.reindex(
        index=range(max_rows),
        columns=range(max_cols),
        fill_value="",
    )
    cmp_df = cmp_df.reindex(
        index=range(max_rows),
        columns=range(max_cols),
        fill_value="",
    )

    ref_norm = ref_df.apply(lambda c: _normalize_col(c, opts))
    cmp_norm = cmp_df.apply(lambda c: _normalize_col(c, opts))

    mask = ref_norm != cmp_norm
    rows_idx, cols_idx = np.where(mask.values)

    diffs = [
        {
            "feuille": sheet,
            "ligne": int(r) + 1,
            "colonne": get_column_letter(int(c) + 1),
            "valeur_ref": str(ref_df.iloc[int(r), int(c)]),
            "valeur_cmp": str(cmp_df.iloc[int(r), int(c)]),
        }
        for r, c in zip(rows_idx, cols_idx)
    ]
    return diffs, max_rows * max_cols


# ---------------------------------------------------------------------------
# Key-based comparison (vectorized)
# ---------------------------------------------------------------------------

def _compare_by_key(
    ref_df: pd.DataFrame,
    cmp_df: pd.DataFrame,
    sheet: str,
    opts: CompareOptions,
    key_col_idx: int,
) -> tuple[List[Dict[str, Any]], int]:
    """Key-based comparison: align rows by a chosen column.

    Rows from both files are matched by the unique value in the
    key column before comparison, so inserted or reordered rows
    do not cascade into spurious differences.

    Returns (differences, total_cells_compared).
    """

    def to_indexed(df: pd.DataFrame) -> pd.DataFrame:
        if df.empty:
            return pd.DataFrame()
        df = df.astype(str).fillna("").copy()
        if key_col_idx >= df.shape[1]:
            raise ValueError(
                f"Colonne clé (index {key_col_idx}) hors limites"
                f" pour la feuille « {sheet} »."
            )
        keys = df.iloc[:, key_col_idx].str.strip()
        orig_cols = [
            c for c in range(df.shape[1]) if c != key_col_idx
        ]
        data = df.drop(columns=df.columns[key_col_idx])
        data.columns = pd.Index(orig_cols)
        data.index = keys
        return data

    ref_i = to_indexed(ref_df)
    cmp_i = to_indexed(cmp_df)

    if ref_i.empty and cmp_i.empty:
        return [], 0

    all_keys = sorted(
        set(ref_i.index.astype(str).tolist())
        | set(cmp_i.index.astype(str).tolist())
    )
    all_data_cols = sorted(
        {int(c) for c in ref_i.columns.tolist()}
        | {int(c) for c in cmp_i.columns.tolist()}
    )
    if not all_data_cols:
        return [], 0

    ref_i = ref_i.reindex(
        index=all_keys, columns=all_data_cols, fill_value=""
    )
    cmp_i = cmp_i.reindex(
        index=all_keys, columns=all_data_cols, fill_value=""
    )

    ref_norm = ref_i.apply(lambda c: _normalize_col(c, opts))
    cmp_norm = cmp_i.apply(lambda c: _normalize_col(c, opts))

    mask = ref_norm != cmp_norm
    rows_idx, cols_idx = np.where(mask.values)

    diffs = [
        {
            "feuille": sheet,
            "ligne": str(all_keys[int(r)]),
            "colonne": get_column_letter(
                int(all_data_cols[int(c)]) + 1
            ),
            "valeur_ref": str(ref_i.iloc[int(r), int(c)]),
            "valeur_cmp": str(cmp_i.iloc[int(r), int(c)]),
        }
        for r, c in zip(rows_idx, cols_idx)
    ]
    return diffs, len(all_keys) * len(all_data_cols)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def compare_workbooks(
    ref_file,
    cmp_file,
    opts: Optional[CompareOptions] = None,
) -> CompareResult:
    """Compare two Excel workbooks and return a CompareResult."""
    if opts is None:
        opts = CompareOptions()

    ref_sheets = _read_workbook(ref_file)
    cmp_sheets = _read_workbook(cmp_file)

    all_sheets = sorted(
        set(list(ref_sheets.keys()) + list(cmp_sheets.keys()))
    )

    key_col_idx = (
        column_index_from_string(opts.key_col) - 1
        if opts.key_col
        else None
    )

    differences: List[Dict[str, Any]] = []
    sheet_stats: Dict[str, int] = {}

    for sheet in all_sheets:
        ref_df = ref_sheets.get(sheet, pd.DataFrame())
        cmp_df = cmp_sheets.get(sheet, pd.DataFrame())

        if key_col_idx is not None:
            diffs, total = _compare_by_key(
                ref_df, cmp_df, sheet, opts, key_col_idx
            )
        else:
            diffs, total = _compare_positional(
                ref_df, cmp_df, sheet, opts
            )

        differences.extend(diffs)
        sheet_stats[sheet] = total

    return CompareResult(
        differences=differences, sheet_stats=sheet_stats
    )


def differences_to_summary(result: CompareResult) -> pd.DataFrame:
    """Return a per-sheet summary DataFrame."""
    counts = Counter(d["feuille"] for d in result.differences)
    rows = []
    for sheet, total in result.sheet_stats.items():
        n_diff = counts.get(sheet, 0)
        pct = round(100 * n_diff / total, 1) if total > 0 else 0.0
        rows.append({
            "Feuille": sheet,
            "Cellules comparées": total,
            "Différences": n_diff,
            "% modifié": f"{pct} %",
        })
    return pd.DataFrame(rows)


# ---------------------------------------------------------------------------
# Export helpers
# ---------------------------------------------------------------------------

def differences_to_dataframe(
    differences: List[Dict[str, Any]]
) -> pd.DataFrame:
    """Convert the list of differences to a pandas DataFrame."""
    if not differences:
        return pd.DataFrame(
            columns=[
                "feuille", "ligne", "colonne",
                "valeur_ref", "valeur_cmp",
            ]
        )
    return pd.DataFrame(differences)


def differences_to_xml(differences: List[Dict[str, Any]]) -> bytes:
    """Serialise differences to an XML byte string for DB injection."""
    root = ET.Element("differences")

    for diff in differences:
        elem = ET.SubElement(root, "difference")
        ET.SubElement(elem, "feuille").text = str(diff["feuille"])
        ET.SubElement(elem, "ligne").text = str(diff["ligne"])
        ET.SubElement(elem, "colonne").text = str(diff["colonne"])
        ET.SubElement(elem, "valeur_ref").text = str(
            diff["valeur_ref"]
        )
        ET.SubElement(elem, "valeur_cmp").text = str(
            diff["valeur_cmp"]
        )

    raw = ET.tostring(root, encoding="unicode")
    pretty = minidom.parseString(raw).toprettyxml(
        indent="  ", encoding="UTF-8"
    )
    return pretty


def differences_to_xlsx(differences: List[Dict[str, Any]]) -> bytes:
    """Serialise differences to an XLSX byte string."""
    df = differences_to_dataframe(differences)
    df.columns = [
        "Feuille", "Ligne / Clé", "Colonne",
        "Valeur référence", "Valeur comparée",
    ]

    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as writer:
        df.to_excel(
            writer, index=False, sheet_name="Différences"
        )

        worksheet = writer.sheets["Différences"]
        for col in worksheet.columns:
            max_length = max(
                len(str(cell.value or "")) for cell in col
            )
            col_letter = col[0].column_letter
            worksheet.column_dimensions[col_letter].width = (
                max_length + 4
            )

    return buf.getvalue()


def differences_to_docx(differences: List[Dict[str, Any]]) -> bytes:
    """Serialise differences to a DOCX byte string."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    title = doc.add_heading(
        "Rapport de comparaison XLcomparator", level=1
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if not differences:
        doc.add_paragraph(
            "Aucune différence trouvée entre les deux fichiers."
        )
    else:
        doc.add_paragraph(
            f"Nombre total de différences : {len(differences)}"
        )
        doc.add_paragraph("")

        headers = [
            "Feuille", "Ligne / Clé", "Colonne",
            "Valeur référence", "Valeur comparée",
        ]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"

        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            run = hdr_cells[i].paragraphs[0].runs[0]
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            para_fmt = hdr_cells[i].paragraphs[0].paragraph_format
            para_fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            tc = hdr_cells[i]._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "2E4057")
            tcPr.append(shd)

        for diff in differences:
            row_cells = table.add_row().cells
            row_cells[0].text = str(diff["feuille"])
            row_cells[1].text = str(diff["ligne"])
            row_cells[2].text = str(diff["colonne"])
            row_cells[3].text = str(diff["valeur_ref"])
            row_cells[4].text = str(diff["valeur_cmp"])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
